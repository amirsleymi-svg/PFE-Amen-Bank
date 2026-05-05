import os
import logging
import re
from typing import List, Dict, Any, Optional
from dataclasses import dataclass
import pdfplumber
from docx import Document as DocxDocument
from langchain.text_splitter import RecursiveCharacterTextSplitter
from rag.config import settings

logger = logging.getLogger("DocumentProcessor")

@dataclass
class Document:
    content: str
    metadata: Dict[str, Any]

class DocumentProcessor:
    def __init__(self):
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=settings.docs.chunk_size,
            chunk_overlap=settings.docs.overlap,
            length_function=len,
            is_separator_regex=False,
        )

    def clean_text(self, text: str) -> str:
        """Clean text from noise, normalize whitespace and fix encodings."""
        if not text:
            return ""
        
        # Remove multiple newlines/whitespaces
        text = re.sub(r'\s+', ' ', text)
        
        # Basic cleanup for Arabic/French special characters if needed
        # (Though usually standard libraries handle this, we ensure it's clean)
        text = text.strip()
        
        return text

    def load_pdf(self, file_path: str) -> str:
        text = ""
        try:
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
        except Exception as e:
            logger.error(f"Error loading PDF {file_path}: {e}")
        return text

    def load_docx(self, file_path: str) -> str:
        text = ""
        try:
            doc = DocxDocument(file_path)
            for para in doc.paragraphs:
                text += para.text + "\n"
        except Exception as e:
            logger.error(f"Error loading DOCX {file_path}: {e}")
        return text

    def load_txt(self, file_path: str) -> str:
        try:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                return f.read()
        except Exception as e:
            logger.error(f"Error loading TXT {file_path}: {e}")
            return ""

    def load_html(self, file_path: str) -> str:
        try:
            from bs4 import BeautifulSoup
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                soup = BeautifulSoup(f.read(), 'html.parser')
                # Remove script and style elements
                for script_or_style in soup(["script", "style"]):
                    script_or_style.decompose()
                return soup.get_text(separator=' ')
        except Exception as e:
            logger.error(f"Error loading HTML {file_path}: {e}")
            return ""

    def process_file(self, file_path: str) -> List[Document]:
        """Loads a file and returns a list of chunked Document objects."""
        ext = os.path.splitext(file_path)[1].lower()
        content = ""

        if ext == '.pdf':
            content = self.load_pdf(file_path)
        elif ext == '.docx':
            content = self.load_docx(file_path)
        elif ext == '.txt':
            content = self.load_txt(file_path)
        elif ext in ['.html', '.htm']:
            content = self.load_html(file_path)
        else:
            logger.warning(f"Unsupported file format: {ext}")
            return []

        if not content:
            logger.warning(f"No content extracted from {file_path}")
            return []

        cleaned_content = self.clean_text(content)
        chunks = self.text_splitter.split_text(cleaned_content)
        
        documents = []
        for i, chunk in enumerate(chunks):
            if len(chunk) >= settings.docs.min_chunk_size:
                documents.append(Document(
                    content=chunk,
                    metadata={
                        "source": os.path.basename(file_path),
                        "chunk_id": i,
                        "file_path": file_path
                    }
                ))
        
        logger.info(f"Processed {file_path}: {len(documents)} chunks created.")
        return documents
